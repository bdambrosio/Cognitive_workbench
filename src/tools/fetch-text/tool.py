#!/usr/bin/env python3
"""
Fetch text from URL or base64 PDF. Auto-detects format (PDF/HTML/MD/TXT) and extracts all text.
"""
import base64
import json
import re
import logging
import tempfile
import os
from urllib.request import urlopen, Request
from urllib.error import URLError, HTTPError
from urllib.parse import urljoin, urlparse
from html.parser import HTMLParser
from typing import Any, Dict, List, Optional
import pymupdf
import warnings
# InfospaceExecutor was removed with the OODA cleanup. The runtime
# `executor` argument is now duck-typed: callers (e.g. chat_loop's
# _FetchTextStubExecutor) just need a _create_uniform_return method.
# Annotations widened to Any to keep the signatures honest about that.

from utils.grobid import parse_pdf_grobid

try:
    from unstructured.partition.html import partition_html
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False

try:
    from playwright.sync_api import sync_playwright
    HAS_PLAYWRIGHT = True
except ImportError:
    HAS_PLAYWRIGHT = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)
warnings.filterwarnings('ignore', category=UserWarning, module='unstructured')


def _fail(executor: Any, reason: str, value: str | None = None, extra: dict | None = None):
    return executor._create_uniform_return(
        "failed",
        value=value or reason,
        reason=reason,
        extra=extra,
    )


def _success(executor: Any, result: str, extra: dict | None = None):
    return executor._create_uniform_return("success", value=result, extra=extra)


_REACT_OBS_CAP = 8000  # cap fetched text in the ReAct working log

# GROBID gives research PDFs a section structure pymupdf cannot; it is
# optional, and any failure falls back to pymupdf's flat page text.
_GROBID_DEFAULT_URL = "http://localhost:8070/api/processFulltextDocument"

# Inline the abstract with the section index when it is small enough to be
# free — it is what the caller almost always wants first, and including it
# saves a whole extra round trip.
_INLINE_ABSTRACT_CAP = 2500


def _resolve_section(requested: str, names: List[str]) -> Optional[str]:
    """Map a requested section name onto one from the index: exact, then
    case-insensitive, then an unambiguous case-insensitive prefix. Returns
    None when nothing matches or a prefix is ambiguous — this resolves an
    identifier the caller read off our own index, so it stays literal
    rather than guessing at intent."""
    req = (requested or '').strip()
    if not req:
        return None
    if req in names:
        return req
    low = req.lower()
    ci = [n for n in names if n.lower() == low]
    if len(ci) == 1:
        return ci[0]
    pref = [n for n in names if n.lower().startswith(low)]
    return pref[0] if len(pref) == 1 else None


def react_invoke(args, *, character_name=None, backend=None, logger=None):
    """ReAct entry-point — see Skill.md for the args contract."""
    from utils.chat_tool_stub import build_tool_kwargs, CapturingResourceManager, translate_result
    url = args.get("url", "")
    if not isinstance(url, str) or not url.strip():
        return {"status": "error", "text": "fetch-text requires non-empty `url`"}

    mgr = CapturingResourceManager()
    result = tool(url, **build_tool_kwargs(
        character_name=character_name, backend=backend, manager=mgr,
    ))

    # fetch-text stashes the extracted text in extra.text.
    out = translate_result(result, manager=mgr, text_key="extra.text",
                           empty_text=f"no text extracted from {url}")

    # Structured provenance for the document this harness actually read.
    # search-web records the sources a model *named*; until this landed
    # nothing recorded a source anyone here opened, so a justify trail
    # could not distinguish a read document from a summarised one — every
    # source in the corpus arrived via search-web. Attached before the
    # section branches below so all ok paths carry it.
    if out.get("status") == "ok":
        _extra = (result.get("extra") or {}) if isinstance(result, dict) else {}
        _md = _extra.get("metadata") or {}
        _title = ((_md.get("pdf_metadata") or {}).get("title")
                  or (_md.get("html_metadata") or {}).get("title") or "")
        try:
            _domain = urlparse(url).netloc.lower().removeprefix("www.")
        except ValueError as e:
            logger.warning(f"fetch-text: cannot parse domain from {url!r}: {e}")
            _domain = ""
        out["meta"] = [{"source_skill": "fetch-text",
                        "tool_metadata": {"sources": [
                            {"url": url, "domain": _domain,
                             "title": _title, "excerpt": ""}]}}]

    # Section-addressable path: available when GROBID parsed the PDF into
    # a section index (research papers). Everything else falls through to
    # the flat capped text below.
    extra = (result.get("extra") or {}) if isinstance(result, dict) else {}
    sections = extra.get("sections") or []
    section_text = extra.get("section_text") or {}
    if out["status"] == "ok" and sections:
        names = [s["name"] for s in sections]
        requested = args.get("section")
        if isinstance(requested, str) and requested.strip():
            name = _resolve_section(requested, names)
            if name is None:
                return {"status": "error",
                        "text": (f"no section named {requested!r} in this "
                                 f"document. Available sections: "
                                 + ", ".join(names))}
            body = section_text.get(name, "")
            if len(body) > _REACT_OBS_CAP:
                body = (body[:_REACT_OBS_CAP].rstrip()
                        + f"\n…[section truncated at {_REACT_OBS_CAP} chars]")
            out["text"] = f"## {name}\n\n{body}"
            return out
        title = ((extra.get("metadata") or {}).get("pdf_metadata") or {}).get("title") or ""
        total = sum(s["chars"] for s in sections)
        lines = [f"# {title}" if title else "# (untitled document)", ""]
        abstract = next((section_text.get(n, "") for n in names
                         if n.strip().lower() == "abstract"), "")
        if abstract and len(abstract) <= _INLINE_ABSTRACT_CAP:
            lines += ["## Abstract", "", abstract, ""]
        lines.append(f"Section index ({len(sections)} sections, "
                     f"{total:,} chars of body text — not shown). Request "
                     f"one with the same url plus `\"section\": \"<name>\"`:")
        for s in sections:
            lines.append(f"  - {s['name']}  ({s['chars']:,} chars)")
        out["text"] = "\n".join(lines)
        return out

    if out["status"] == "ok" and len(out["text"]) > _REACT_OBS_CAP:
        out["text"] = out["text"][:_REACT_OBS_CAP].rstrip() + \
            f"\n…[truncated at {_REACT_OBS_CAP} chars]"

    # Surface HTML image candidates. Page text alone doesn't contain
    # og:image meta-tag URLs, so without this a downstream `display`
    # call has to re-fetch and re-parse to pick a primary image.
    # Appended after truncation so the candidates always survive.
    if out["status"] == "ok":
        images = (result.get("extra") or {}).get("metadata", {}) \
            .get("html_metadata", {}).get("images", {}) or {}
        lines = []
        if images.get("og_image"):
            lines.append(f"- og:image: {images['og_image']}")
        if images.get("twitter_image"):
            lines.append(f"- twitter:image: {images['twitter_image']}")
        for img in (images.get("imgs") or [])[:3]:
            src = img.get("src", "")
            alt = img.get("alt", "")
            if src:
                lines.append(f"- <img>: {src}" + (f"  (alt: {alt})" if alt else ""))
        if lines:
            out["text"] = out["text"] + "\n\nImage candidates:\n" + "\n".join(lines)

    return out


def tool(url_or_content: str, runtime=None, **kwargs) -> str:
    """
    Fetch text from URL, base64 PDF, Note ID, or Collection ID. Auto-detects format and extracts all text.
    Collection-aware: If input is Collection ID, fetches first item's content.
    Note-aware: If input is Note ID, retrieves Note's content directly.
    
    Args:
        url_or_content: URL string, base64-encoded PDF content, Note ID, or Collection ID
        **kwargs: grobid_url (optional) - GROBID server URL for PDF parsing if available in config
        
    Returns:
        JSON string with text, format, metadata, page_count (if PDF), char_count
    """
    executor: Any = kwargs.get("executor")
    if not executor:
        return {"status": "failed", "reason": "executor not available", "value": None, "resource_id": None}

    # Extract grobid_url and pdf_parser from kwargs (from YAML config).
    # build_tool_kwargs does not carry grobid_url, so the chat path always
    # arrived here with None and every PDF silently took the pymupdf
    # branch. Fall back to the environment: GROBID_URL unset means the
    # standard local endpoint, GROBID_URL="" disables GROBID outright.
    # An unreachable server just raises and falls through to pymupdf.
    grobid_url = kwargs.get('grobid_url')
    if grobid_url is None:
        grobid_url = os.environ.get('GROBID_URL', _GROBID_DEFAULT_URL)
    pdf_parser = kwargs.get('pdf_parser')
    
    if not url_or_content or not isinstance(url_or_content, str):
        return _fail(executor, "url_or_content parameter required")
    
    # Check if input is Note ID - retrieve Note content directly
    if url_or_content.startswith('Note_'):
        try:
            resource_mgr = kwargs.get('resource_manager')
            
            # Use resource manager to get Note
            if resource_mgr:
                note_content = resource_mgr.get_resource(url_or_content)
                if note_content:
                    content = note_content.get('properties', {}).get('content', '')
                else:
                    return _fail(executor, f"Note {url_or_content} not found")
            else:
                return _fail(executor, "resource_manager not available")
            
            # Text-only: content is string, return as-is
            text = str(content) if content is not None else ""
            return _success(executor, text)
        except Exception as e:
            logger.error(f"Failed to retrieve Note {url_or_content}: {e}")
            return _fail(executor, "Failed to retrieve Note", extra={"exception": str(e)})
    
    # Check if input is Collection ID - extract first item
    if url_or_content.startswith('Collection_'):
        try:
            resource_mgr = kwargs.get('resource_manager')
            note_ids = []
            
            # Use resource_manager to get Collection
            if resource_mgr:
                collection_resource = resource_mgr.get_resource(url_or_content)
                if collection_resource:
                    note_ids = collection_resource.get('properties', {}).get('content', [])
            else:
                return _fail(executor, "resource_manager not available")
            
            if not note_ids:
                return _fail(executor, f"Collection {url_or_content} not found or empty")
            
            if len(note_ids) == 0:
                return _fail(executor, "Collection is empty")
            
            if len(note_ids) > 1:
                logger.warning(f"fetch-text: Collection {url_or_content} has {len(note_ids)} items, using first")
            
            first_note_id = note_ids[0]
            # Recursively call fetch-text with Note ID (handles both structured and plain Notes)
            return tool(first_note_id, **kwargs)
        except Exception as e:
            logger.error(f"Failed to resolve Collection: {e}")
            return _fail(executor, "Failed to resolve Collection", extra={"exception": str(e)})
    
    # Check if input is local absolute file path
    if url_or_content.startswith('/') and os.path.isfile(url_or_content):
        try:
            with open(url_or_content, 'rb') as f:
                content = f.read()
            content_type = None
            final_url = url_or_content
        except OSError as e:
            logger.error(f"File read failed: {str(e)}")
            return _fail(executor, f"Failed to read file {url_or_content}: {str(e)}")
    else:
        # Check if input is base64 PDF (backward compatibility)
        if _is_base64_pdf(url_or_content):
            payload = _process_base64_pdf(url_or_content)
            return _success(executor, payload)
        
        # Otherwise treat as URL
        url = _extract_url(url_or_content)
        if not url:
            return _fail(executor, f"Invalid URL or content: {url_or_content[:100]}")
        
        # Download and detect format
        content, content_type, final_url, dl_reason = _download_from_url(url)
        if not content:
            return _fail(executor, f"Failed to download from {url}: {dl_reason or 'unknown'}")
    
    # Now detect format and extract (common for both file and URL)
    file_format = _detect_format(content, content_type, final_url)
    
    if file_format == "pdf":
        result = _extract_pdf_text(content, final_url, grobid_url=grobid_url, pdf_parser=pdf_parser)
    elif file_format == "html":
        result = _extract_html_text(content, final_url)
    elif file_format == "docx":
        result = _extract_docx_text(content, final_url)
    elif file_format == "markdown":
        result = _extract_text_plain(content, final_url, "markdown")
    else:
        result = _extract_text_plain(content, final_url, "text")

    # Attempt to include parsed JSON as data for convenience
    extra_data = None
    try:
        extra_data = json.loads(result)
    except Exception:
        extra_data = {"raw_result": result}
    
    if isinstance(extra_data, dict) and "error" in extra_data:
        return _fail(executor, extra_data["error"], value=result, extra=extra_data)
    
    return _success(executor, result, extra_data)


def _is_base64_pdf(content: str) -> bool:
    """Check if content is base64-encoded PDF."""
    if len(content) < 100:
        return False
    try:
        decoded = base64.b64decode(content[:200])
        return decoded.startswith(b'%PDF')
    except:
        return False


def _extract_url(url_str: str) -> str:
    """Extract URL from string, handling JSON wrappers and embedded URLs."""
    extracted_url = url_str.strip()
    
    # Try JSON parsing
    if extracted_url.startswith('{') or extracted_url.startswith('['):
        try:
            parsed = json.loads(extracted_url)
            if isinstance(parsed, dict):
                if 'url' in parsed:
                    return parsed['url']
                # Recursive search
                def find_url(obj):
                    if isinstance(obj, dict):
                        if 'url' in obj:
                            return obj['url']
                        for v in obj.values():
                            result = find_url(v)
                            if result:
                                return result
                    elif isinstance(obj, list):
                        for item in obj:
                            result = find_url(item)
                            if result:
                                return result
                    return None
                found = find_url(parsed)
                if found:
                    return found
        except (json.JSONDecodeError, TypeError):
            pass
    
    # Regex fallback
    if not extracted_url.startswith(('http://', 'https://')):
        url_match = re.search(r'https?://[^\s<>"\'{}|\\^`\[\]]+', extracted_url)
        if url_match:
            extracted_url = url_match.group(0)
    
    # Validate
    if not extracted_url.startswith(('http://', 'https://')):
        return None
    
    return _normalize_url(extracted_url)


def _normalize_url(url: str) -> str:
    """Normalize known URL patterns."""
    # ArXiv: convert /abs/ to /pdf/
    if 'arxiv.org/abs/' in url:
        url = url.replace('/abs/', '/pdf/')
        if not url.endswith('.pdf'):
            url = url + '.pdf'
    # GitHub: convert /blob/ to /raw/ to get raw content instead of HTML
    if 'github.com' in url and '/blob/' in url:
        url = url.replace('/blob/', '/raw/')
    return url


def _download_from_url(url: str) -> tuple:
    """Download content from URL. Returns (content_bytes, content_type,
    final_url, reason). On success, reason is None; on failure, content
    and content_type are None and reason is a short string the caller
    can surface to the LLM observation. The error reason matters for
    ReAct: a generic "failed to download" prevents the model from
    judging whether to retry differently (e.g., URLError on basic-auth-
    in-URL → caller might extract creds and inject a header instead)."""
    headers = {'User-Agent': 'Mozilla/5.0 (compatible; CognitiveWorkbench/1.0)'}

    # Add HuggingFace authentication if URL is from HuggingFace and token is available
    if 'huggingface.co' in url:
        hf_token = os.getenv('HF_TOKEN') or os.getenv('HUGGINGFACE_TOKEN')
        if hf_token:
            headers['Authorization'] = f'Bearer {hf_token}'

    try:
        request = Request(url, headers=headers)
        response = urlopen(request, timeout=30)
        content = response.read()
        content_type = response.headers.get('Content-Type', '').lower()
        final_url = response.url or url

        return content, content_type, final_url, None
    except HTTPError as e:
        reason = f"HTTP {e.code} {e.reason}"
        logger.error(f"{reason} for {url}")
        return None, None, url, reason
    except URLError as e:
        reason = f"URLError: {e.reason}"
        logger.error(f"{reason} for {url}")
        return None, None, url, reason
    except Exception as e:
        reason = f"{type(e).__name__}: {e}"
        logger.error(f"Download failed: {reason} for {url}")
        return None, None, url, reason


def _detect_format(content: bytes, content_type: str, url: str) -> str:
    """Detect file format from content, Content-Type header, and URL."""
    # Check magic bytes first (most reliable)
    if content.startswith(b'%PDF'):
        return "pdf"
    # DOCX files are ZIP archives (Office Open XML)
    if content.startswith(b'PK') and len(content) > 4:
        url_lower = url.lower()
        # Check for DOCX (Microsoft Word)
        if url_lower.endswith('.docx') or 'msword' in content_type or 'wordprocessingml' in content_type:
            return "docx"
    
    # Check Content-Type header (more reliable than URL extension)
    if 'pdf' in content_type:
        return "pdf"
    if 'msword' in content_type or 'wordprocessingml' in content_type:
        return "docx"
    if 'markdown' in content_type:
        return "markdown"
    if 'html' in content_type:
        return "html"
    
    # Check content for HTML markers (override URL extension if HTML detected)
    content_start = content[:1024].decode('utf-8', errors='ignore').strip()
    if content_start.startswith('<!DOCTYPE html>') or content_start.startswith('<!doctype html>'):
        return "html"
    if content.startswith(b'<') or b'<html' in content[:1024].lower():
        return "html"
    
    # Check URL extension (least reliable - only if content didn't indicate HTML)
    url_lower = url.lower()
    if url_lower.endswith('.docx'):
        return "docx"
    if url_lower.endswith('.md') or url_lower.endswith('.markdown'):
        return "markdown"
    if url_lower.endswith('.html') or url_lower.endswith('.htm'):
        return "html"
    if url_lower.endswith('.txt'):
        return "text"
    
    # Check text/plain with markdown detection
    if 'text/plain' in content_type:
        try:
            text_sample = content[:2048].decode('utf-8', errors='ignore')
            if _looks_like_markdown(text_sample):
                return "markdown"
        except:
            pass
        return "text"
    
    # Check content for markdown patterns
    try:
        text_sample = content[:2048].decode('utf-8', errors='ignore')
        if _looks_like_markdown(text_sample):
            return "markdown"
    except:
        pass
    
    # Default to text
    return "text"


def _looks_like_markdown(text: str) -> bool:
    """Heuristic check if text looks like markdown."""
    if not text or len(text) < 20:
        return False
    
    # Check for markdown patterns
    markdown_indicators = [
        r'^#+\s',  # Headers
        r'^\*\s',  # Unordered lists
        r'^\d+\.\s',  # Ordered lists
        r'\[.*?\]\(.*?\)',  # Links
        r'```',  # Code fences
        r'\*\*.*?\*\*',  # Bold
        r'_.*?_',  # Italic
    ]
    
    lines = text.split('\n')[:50]  # Check first 50 lines
    matches = 0
    for line in lines:
        for pattern in markdown_indicators:
            if re.search(pattern, line):
                matches += 1
                break
    
    # If 10%+ of lines match markdown patterns, likely markdown
    return matches >= max(2, len(lines) * 0.1)


def _process_base64_pdf(content: str) -> str:
    """Process base64 PDF content (backward compatibility)."""
    try:
        pdf_bytes = base64.b64decode(content)
        return _extract_pdf_text(pdf_bytes, "")
    except Exception as e:
        return json.dumps({"error": f"Failed to decode PDF: {str(e)}"})


def _extract_pdf_text(content: bytes, url: str, grobid_url: str = None, pdf_parser: str = None) -> str:
    """
    Extract text from PDF bytes.
    
    Args:
        content: PDF file bytes
        url: Source URL or file path
        grobid_url: Optional GROBID server URL for enhanced PDF parsing
        pdf_parser: When "pymupdf", use pymupdf only (skip GROBID even if available)
    """
    use_pymupdf_only = (pdf_parser or "").lower() == "pymupdf"
    # Use GROBID if available and not forcing pymupdf
    if grobid_url and not use_pymupdf_only:
        try:
            # Save content to temp file for GROBID
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(content)
                tmp_filepath = tmp_file.name
            
            # Extract title from URL for temp filename
            title = url.split('/')[-1].replace('.pdf', '') or "document"
            
            # Parse with GROBID
            grobid_result = parse_pdf_grobid(pdf_filepath=tmp_filepath, title=title, grobid_url=grobid_url)
            
            # Clean up temp file
            try:
                os.unlink(tmp_filepath)
            except:
                pass
            
            if grobid_result:
                # Concatenate chunks with section headers
                chunks = grobid_result.get('chunks', [])
                # Group chunks under their section heading, preserving
                # document order. This index is what makes a paper
                # addressable: the caller reads the section list first and
                # asks for the one or two sections it needs, instead of
                # pulling the whole body (a long paper runs to ~40k tokens)
                # into a context it has to keep re-sending.
                ordered: List[str] = []
                by_section: Dict[str, List[str]] = {}
                for section_title, section_text in chunks:
                    name = str(section_title or '').strip() or 'Untitled section'
                    if name not in by_section:
                        by_section[name] = []
                        ordered.append(name)
                    by_section[name].append(section_text or '')
                section_text_map = {
                    n: "\n\n".join(by_section[n]).strip() for n in ordered}

                if chunks:
                    chunk_texts = []
                    for section_title, section_text in chunks:
                        chunk_texts.append(f"{section_title}\n{section_text}")
                    full_text = "\n\n".join(chunk_texts)
                elif grobid_result.get('abstract'):
                    full_text = grobid_result['abstract']
                else:
                    full_text = ""

                pdf_title = (grobid_result.get("title") or "").strip()
                if pdf_title:
                    full_text = f"# {pdf_title}\n\n{full_text}"

                result = {
                    "text": full_text,
                    "format": "pdf",
                    "metadata": {
                        "source_url": url,
                        "pdf_metadata": {
                            "title": pdf_title,
                            "author": grobid_result.get("authors", ""),
                            "subject": "",
                            "creator": ""
                        }
                    },
                    "sections": [{"name": n, "chars": len(section_text_map[n])}
                                 for n in ordered],
                    "section_text": section_text_map,
                    "page_count": len(chunks) if chunks else 0,
                    "char_count": len(full_text)
                }

                logger.info(f"Extracted {len(full_text)} chars from PDF using GROBID ({len(chunks)} chunks)")
                return json.dumps(result, indent=2)
        except Exception as e:
            logger.warning(f"GROBID parsing failed, falling back to pymupdf: {str(e)}")
            # Fall through to pymupdf extraction
    
    # Fallback to pymupdf extraction
    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
        pages_text = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            pages_text.append(text)
        
        full_text = "\n\n".join(pages_text)
        metadata = doc.metadata or {}
        page_count = len(doc)
        doc.close()

        pdf_title = (metadata.get("title") or "").strip()
        if pdf_title:
            full_text = f"# {pdf_title}\n\n{full_text}"

        result = {
            "text": full_text,
            "format": "pdf",
            "metadata": {
                "source_url": url,
                "pdf_metadata": {
                    "title": pdf_title,
                    "author": metadata.get("author", ""),
                    "subject": metadata.get("subject", ""),
                    "creator": metadata.get("creator", "")
                }
            },
            "page_count": page_count,
            "char_count": len(full_text)
        }

        logger.info(f"Extracted {len(full_text)} chars from {page_count} page PDF using pymupdf")
        return json.dumps(result, indent=2)
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        return json.dumps({"error": f"PDF extraction failed: {str(e)}"})


def _extract_docx_text(content: bytes, url: str) -> str:
    """Extract text from DOCX content."""
    if not HAS_DOCX:
        return json.dumps({"error": "python-docx library not installed"})
    
    try:
        # DOCX files are ZIP archives, need to write to temp file for python-docx
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(content)
            tmp_path = tmp_file.name
        
        try:
            doc = Document(tmp_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            full_text = '\n'.join(paragraphs)
            
            # Extract metadata
            core_props = doc.core_properties
            docx_metadata = {}
            if core_props.title:
                docx_metadata['title'] = core_props.title
            if core_props.author:
                docx_metadata['author'] = core_props.author
            if core_props.subject:
                docx_metadata['subject'] = core_props.subject
            
            result = {
                "text": full_text,
                "format": "docx",
                "metadata": {
                    "source_url": url,
                    "docx_metadata": docx_metadata
                },
                "char_count": len(full_text)
            }
            
            return json.dumps(result, ensure_ascii=False)
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return json.dumps({"error": f"DOCX extraction failed: {str(e)}"})


def _extract_html_text(content: bytes, url: str) -> str:
    """Extract text from HTML content."""
    html_str = content.decode('utf-8', errors='ignore')
    extraction_method = "static_html"
    
    # Try normal extraction first
    if HAS_UNSTRUCTURED:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            elements = partition_html(text=html_str)
        text_parts = [str(e).strip() for e in elements if str(e).strip()]
        full_text = "\n".join(text_parts)
    else:
        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text = []
            def handle_data(self, data):
                if data.strip():
                    self.text.append(data.strip())
        parser = TextExtractor()
        parser.feed(html_str)
        full_text = "\n".join(parser.text)
    
    # If we got very little text, try playwright fallback (for SPAs)
    if len(full_text) < 100 and HAS_PLAYWRIGHT:
        logger.info(f"Low text content ({len(full_text)} chars) from {url}, trying playwright fallback")
        playwright_text = _try_playwright_extraction(url)
        if playwright_text and len(playwright_text) > len(full_text):
            full_text = playwright_text
            extraction_method = "playwright"
            logger.info(f"Playwright extraction succeeded: {len(full_text)} chars")
    
    # Extract basic HTML metadata
    html_metadata = {}
    title_match = re.search(r'<title[^>]*>(.*?)</title>', html_str, re.IGNORECASE | re.DOTALL)
    if title_match:
        html_metadata["title"] = re.sub(r'<[^>]+>', '', title_match.group(1)).strip()

    # Image candidates — surfaced so a downstream consumer (e.g. the
    # chat-loop `display` tool) can pick a "primary" image semantically
    # without having to refetch and reparse the page. Resolved against
    # the final URL so relative srcs are usable as-is.
    page_images = _extract_page_images(html_str, url)
    if page_images:
        html_metadata["images"] = page_images
    
    result = {
        "text": full_text,
        "format": "html",
        "metadata": {
            "source_url": url,
            "html_metadata": html_metadata,
            "extraction_method": extraction_method
        },
        "char_count": len(full_text)
    }
    
    logger.info(f"Extracted {len(full_text)} chars from HTML using {extraction_method}")
    return json.dumps(result, indent=2)


def _extract_page_images(html_str: str, page_url: str, max_imgs: int = 5) -> dict:
    """Pull image candidates out of an HTML page.

    Returns a dict with any of: og_image, twitter_image, imgs (list of
    {src, alt}). Empty dict if the page has no images. All srcs are
    absolute (resolved against page_url) so consumers can pass them
    straight to the canvas proxy without re-resolving.

    Capped at `max_imgs` for the imgs list so we don't drag the whole
    page into the result — Jill needs candidates to choose from, not an
    exhaustive index.
    """
    out: dict = {}

    def _resolve(u: str) -> str:
        u = (u or '').strip()
        if not u:
            return ''
        try:
            return urljoin(page_url, u)
        except Exception as e:
            logger.debug(f"urljoin failed for {u!r} against {page_url!r}: {e}")
            return u

    def _meta(prop_or_name: str, key: str) -> str:
        # Tolerate either attribute order: meta tags routinely appear as
        # both `<meta property="og:image" content="...">` and the reverse.
        esc = re.escape(prop_or_name)
        m = re.search(
            rf'<meta[^>]+\b{key}=["\']{esc}["\'][^>]+content=["\']([^"\']+)["\']',
            html_str, re.IGNORECASE,
        ) or re.search(
            rf'<meta[^>]+content=["\']([^"\']+)["\'][^>]+\b{key}=["\']{esc}["\']',
            html_str, re.IGNORECASE,
        )
        return _resolve(m.group(1)) if m else ''

    og = _meta('og:image', 'property')
    tw = _meta('twitter:image', 'name')
    if og:
        out['og_image'] = og
    if tw and tw != og:
        out['twitter_image'] = tw

    imgs: list = []
    seen: set = set()
    for m in re.finditer(r'<img\b([^>]*?)>', html_str, re.IGNORECASE):
        if len(imgs) >= max_imgs:
            break
        attrs = m.group(1)
        src_match = re.search(r'\bsrc=["\']([^"\']+)["\']', attrs, re.IGNORECASE)
        if not src_match:
            continue
        src = _resolve(src_match.group(1))
        if not src or src.startswith('data:') or src in seen:
            continue
        alt_match = re.search(r'\balt=["\']([^"\']*)["\']', attrs, re.IGNORECASE)
        alt = alt_match.group(1).strip() if alt_match else ''
        imgs.append({'src': src, 'alt': alt})
        seen.add(src)
    if imgs:
        out['imgs'] = imgs
    return out


def _try_playwright_extraction(url: str) -> str:
    """Try extracting text using playwright (for SPAs)."""
    if not url.startswith(('http://', 'https://')):
        logger.info(f"Skipping playwright for non-URL: {url}")
        return ""

    if not HAS_PLAYWRIGHT:
        return ""

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, wait_until='networkidle', timeout=30000)
            # Use inner_text to get rendered text content, not HTML markup
            full_text = page.inner_text('body')
            browser.close()
            return full_text
    except Exception as e:
        logger.warning(f"Playwright extraction failed for {url}: {str(e)}")
        return ""


def _extract_text_plain(content: bytes, url: str, format_type: str) -> str:
    """Extract text from plain text or markdown content."""
    try:
        text = content.decode('utf-8', errors='ignore')
        
        result = {
            "text": text,
            "format": format_type,
            "metadata": {
                "source_url": url
            },
            "char_count": len(text)
        }
        
        logger.info(f"Extracted {len(text)} chars from {format_type}")
        return json.dumps(result, indent=2)
    except Exception as e:
        logger.error(f"Text extraction failed: {str(e)}")
        return json.dumps({"error": f"Text extraction failed: {str(e)}"})

